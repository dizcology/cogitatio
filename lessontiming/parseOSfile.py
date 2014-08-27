from docx import *
import re

def getItemNum(string):
    itemno_pat=r'^\d{2,3}'
    res=re.findall(itemno_pat,string)
    if res:
        return res[0].zfill(3)
    else:
        return None
  

def parseOSfile(osfn):
    '''Parse the given OS file, extracting the paths of interest.'''
    osfile = Document(osfn)

    paths = {'weak + behind': [], 
            'weak + ontime': [],
            'branches': []}     # 'branches' is different from the others: it hold
                                # all options for 'unspecified' branches.

    truepaths = ['weak + behind','weak + ontime']

    #itemno_pat='\s*[0-9][0-9][0-9]\s*\.?'  #used to be ' ?[0-9][0-9][0-9]?\.' which makes little sense

    for par in osfile.paragraphs:   
    # First, go through the paragraphs and pull out any numbers starting lines.
        tempnum=getItemNum(par.text)
        if tempnum:
            itemno = tempnum
            if 'quiz' in par.text.lower():
                pass
            elif 'skip if behind' in par.text.lower():
                paths['weak + ontime'].append(itemno.encode('ascii'))
            else:
                for path in truepaths:
                    paths[path].append(itemno.encode('ascii'))

    osfile = Document(osfn)
    for tab in osfile.tables:
    # Next, go through the tables.
        branchpaths = []        # Keep track of "undefined" branches - i.e. those not defined by strength and speed
        if  len(tab.columns)==1:
            col=tab.columns[0]
            for par in col.cells[0].paragraphs+col.cells[1].paragraphs: #temporary fix until everyone uses the new templates of script_drafter
                tempnum = getItemNum(par.text)
                if tempnum:
                    itemno = tempnum
            for par in col.cells[0].paragraphs:
              if 'quiz' in par.text.lower():
                  pass
              elif 'skip if behind' in par.text.lower():
                  paths['weak + ontime'].append(itemno.encode('ascii'))
              else:
                  for path in truepaths:
                      paths[path].append(itemno.encode('ascii'))
            continue
        if len(tab.columns[0].cells) > 1:
            defaultitemno = tab.columns[0].cells[1].paragraphs[0].text.split('.')[0].replace(' ','').zfill(3)
            # defaultitemno is the fallback if the OS describes the branch as 'same as...'.
            for col in tab.columns:
                try:
                    branchpaths.append([])
                    colheader = col.cells[0].paragraphs[0].text
                    for par in col.cells[1].paragraphs:
                        tempnum = getItemNum(par.text)
                        if tempnum:
                            itemno = tempnum
                        #elif re.match('^same',par.text.lower()):   # Removing this because it's not used in the weak
                        #   itemno = defaultitemno                  # but potentially causes problems.
                        else: itemno = ''
                        if not itemno == '':
                            if 'weak' in colheader.lower():
                                if not ('skip if behind' in colheader.lower() or 'not behind' in colheader.lower()):
                                    paths['weak + behind'].append(itemno.encode('ascii'))
                                if not colheader.lower() == 'behind':
                                    paths['weak + ontime'].append(itemno.encode('ascii'))
                            elif not ('average' in colheader.lower() or 'strong' in colheader.lower()):
                                if 'behind' in colheader.lower():
                                    if not ('skip if behind' in colheader.lower() or 'not behind' in colheader.lower()):
                                        paths['weak + behind'].append(itemno.encode('ascii'))
                                    if not colheader.lower() == 'behind':
                                        paths['weak + ontime'].append(itemno.encode('ascii'))
                                else:
                                    branchpaths[-1].append(itemno.encode('ascii'))
                except IndexError:
                    pass
                except Exception as e: 
                    print e
                    pass
            if sum([len(p) for p in branchpaths]) > 0:
                paths['branches'].append(branchpaths)

    for path in truepaths: 
        #paths[path] = sorted(list(set(paths[path])))        # Remove duplicates and sort... but why should there be duplicates?
        while "" in paths[path]:
          paths[path].remove("")
        
        paths[path] = sorted(list(set(paths[path])))
    return paths
